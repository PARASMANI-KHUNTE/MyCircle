
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

def set_heading(para, level, text, color=None):
    style_name = f'Heading {level}'
    if style_name not in [s.name for s in doc.styles]:
        style_name = 'Normal'
    para.style = doc.styles[style_name] if style_name != 'Normal' else doc.styles['Normal']
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16 - level * 2) if level < 4 else Pt(12)
    if color:
        run.font.color.rgb = RGBColor(*color)

def h(level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(max(12, 18 - level * 2))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(text, justify=True):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def center(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    if col_widths:
        for i, row in enumerate(t.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
center("GURU GHASIDAS VISHWAVIDYALAYA, BILASPUR (C.G.)", bold=True, size=16)
center("(A Central University)", bold=False, size=13)
center("Department of Computer Science and Information Technology", bold=True, size=13)
doc.add_paragraph()
center("MINOR PROJECT REPORT", bold=True, size=15)
center("Submitted in partial fulfillment of the requirements", bold=False, size=12)
center("for the award of the degree of", bold=False, size=12)
center("MASTER OF COMPUTER APPLICATIONS", bold=True, size=14)
doc.add_paragraph()
center("On", bold=False, size=12)
doc.add_paragraph()
center("MyCircle – A Local Student Opportunity & Networking Platform", bold=True, size=16)
doc.add_paragraph()
doc.add_paragraph()
center("Submitted By:", bold=True, size=12)
center("[STUDENT_NAME]", bold=False, size=12)
center("Roll No: [ROLL_NO]", bold=False, size=12)
center("Enrollment No: [ENROLLMENT_NO]", bold=False, size=12)
center("MCA 2nd Semester", bold=False, size=12)
doc.add_paragraph()
center("Guided By:", bold=True, size=12)
center("[GUIDE_NAME]", bold=False, size=12)
center("Assistant Professor / Associate Professor", bold=False, size=12)
center("Department of Computer Science and IT", bold=False, size=12)
doc.add_paragraph()
doc.add_paragraph()
center("Academic Year: 2024–2025", bold=True, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CERTIFICATE OF APPROVAL
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CERTIFICATE OF APPROVAL")
doc.add_paragraph()
body('This is to certify that the Minor Project Report entitled "MyCircle – A Local Student Opportunity & Networking Platform" submitted by [STUDENT_NAME], Roll No. [ROLL_NO], Enrollment No. [ENROLLMENT_NO], a student of Master of Computer Applications (MCA), 2nd Semester, Department of Computer Science and Information Technology, Guru Ghasidas Vishwavidyalaya, Bilaspur (C.G.), is a bonafide record of the work carried out by the student under the supervision of [GUIDE_NAME].')
body("The project has been completed in accordance with the academic guidelines prescribed by Guru Ghasidas Vishwavidyalaya and fulfills the requirements for the award of the degree of Master of Computer Applications. The work embodied in this report has not been submitted elsewhere, either in full or in part, for the award of any degree or diploma.")
body("We, the undersigned, hereby approve this Minor Project Report and recommend it for acceptance.")
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run("Internal Examiner:\t\t\t\t\tExternal Examiner:").font.name = 'Times New Roman'
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("_______________________\t\t\t\t_______________________").font.name = 'Times New Roman'
p = doc.add_paragraph()
p.add_run("Name & Designation\t\t\t\t\tName & Designation").font.name = 'Times New Roman'
doc.add_paragraph()
center("Head of Department", bold=True, size=12)
center("Department of Computer Science and IT", bold=False, size=12)
center("Guru Ghasidas Vishwavidyalaya, Bilaspur (C.G.)", bold=False, size=12)
center("Date: _______________", bold=False, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# GUIDE CERTIFICATE
# ══════════════════════════════════════════════════════════════════════════════
h(1, "GUIDE CERTIFICATE")
doc.add_paragraph()
body('This is to certify that the Minor Project work entitled "MyCircle – A Local Student Opportunity & Networking Platform" has been carried out by [STUDENT_NAME] (Roll No: [ROLL_NO], Enrollment No: [ENROLLMENT_NO]), student of MCA 2nd Semester, Department of Computer Science and Information Technology, Guru Ghasidas Vishwavidyalaya, Bilaspur (C.G.), under my direct supervision and guidance.')
body("The project has been completed as per the norms and guidelines laid down by the University. The student has worked diligently and sincerely on the project, and the results presented are genuine, original, and have not been copied or reproduced from any other source. To the best of my knowledge, this work has not been submitted, either wholly or in part, to any other university or institution for the award of any degree or diploma.")
body("I recommend this Minor Project Report for evaluation and approval by the Department and the University.")
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Guide / Supervisor:").bold = True
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("_______________________")
p = doc.add_paragraph()
p.add_run("[GUIDE_NAME]")
p = doc.add_paragraph()
p.add_run("Assistant Professor / Associate Professor")
p = doc.add_paragraph()
p.add_run("Department of Computer Science and IT")
p = doc.add_paragraph()
p.add_run("Guru Ghasidas Vishwavidyalaya, Bilaspur (C.G.)")
p = doc.add_paragraph()
p.add_run("Date: _______________")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SELF DECLARATION
# ══════════════════════════════════════════════════════════════════════════════
h(1, "SELF DECLARATION")
doc.add_paragraph()
body('I, [STUDENT_NAME], Roll No. [ROLL_NO], Enrollment No. [ENROLLMENT_NO], a student of Master of Computer Applications (MCA), 2nd Semester, Department of Computer Science and Information Technology, Guru Ghasidas Vishwavidyalaya, Bilaspur (C.G.), hereby declare that the Minor Project Report entitled "MyCircle – A Local Student Opportunity & Networking Platform" submitted in partial fulfillment of the requirements for the award of the degree of Master of Computer Applications is an authentic record of my own work carried out during the academic year 2024–2025.')
body("I further declare that the information presented in this report is true and correct to the best of my knowledge, and no part of it has been submitted previously, either in full or in part, to any other university, institution, or organization for the award of any degree, diploma, or certificate. All sources of information used in this project have been duly acknowledged and referenced.")
body("I understand that any misrepresentation or plagiarism will be treated as a serious academic offence and may result in the cancellation of my degree.")
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Place: Bilaspur")
p = doc.add_paragraph()
p.add_run("Date: _______________")
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("_______________________")
p = doc.add_paragraph()
p.add_run("[STUDENT_NAME]")
p = doc.add_paragraph()
p.add_run("Roll No: [ROLL_NO]")
p = doc.add_paragraph()
p.add_run("Enrollment No: [ENROLLMENT_NO]")
p = doc.add_paragraph()
p.add_run("MCA 2nd Semester, GGU, Bilaspur")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════════════════════
h(1, "ACKNOWLEDGEMENT")
doc.add_paragraph()
body("The successful completion of this Minor Project Report would not have been possible without the guidance, encouragement, and support of several individuals, to whom I owe my deepest gratitude. I take this opportunity to express my sincere appreciation to all those who have contributed, directly or indirectly, to the realization of this project.")
body("First and foremost, I express my heartfelt gratitude to my project guide, [GUIDE_NAME], Assistant Professor / Associate Professor, Department of Computer Science and Information Technology, Guru Ghasidas Vishwavidyalaya, for their invaluable guidance, continuous encouragement, and scholarly insights throughout the course of this project. Their constructive feedback and patient mentoring helped me navigate the challenges encountered during both the development and documentation phases of MyCircle.")
body("I am also deeply grateful to the Head of the Department of Computer Science and Information Technology, and all the faculty members of the department, for providing an intellectually stimulating environment and equipping me with the academic foundation necessary to undertake a project of this scope. The lectures, seminars, and hands-on laboratory sessions conducted throughout the MCA curriculum were instrumental in building the technical competency required for full-stack web development.")
body("I extend my sincere thanks to my classmates and friends who participated in the User Acceptance Testing (UAT) phase of MyCircle, providing honest feedback that significantly improved the usability and design of the platform. Their enthusiasm and collaborative spirit were a constant source of motivation. Finally, I owe an immeasurable debt of gratitude to my family, whose unwavering support, patience, and encouragement gave me the strength to persevere through the demanding hours of coding, debugging, and writing that this project demanded.")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("[STUDENT_NAME]")
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
h(1, "ABSTRACT")
doc.add_paragraph()
body("MyCircle is a full-stack web application developed using the MERN technology stack—MongoDB, Express.js, React.js, and Node.js—specifically designed to address the chronic lack of a centralized, local opportunity-discovery and peer-networking platform for students enrolled in colleges and universities across smaller Indian cities and towns. Despite the proliferation of mainstream social and professional networks, students at institutions like those affiliated with Guru Ghasidas Vishwavidyalaya, Bilaspur, continue to miss out on internships, freelance gigs, part-time jobs, campus hackathons, study groups, and collaborative academic opportunities, simply because no platform exists that caters to their hyper-local, institution-specific needs.")
body("MyCircle fills this critical gap by providing a dedicated space where students can create verified institutional profiles, post and discover local opportunities across multiple categories (internships, gigs, jobs, barter, and rentals), follow peers, engage in real-time one-on-one messaging, and receive notifications about new activity within their academic circle. The platform incorporates JWT-based authentication for secure session management, bcrypt for password hashing, Cloudinary integration for cloud-based media storage, and Socket.io for real-time event-driven communication. The frontend is developed using React.js with Vite as the build tool, delivering a responsive, mobile-first user interface styled with Tailwind CSS. The backend is a stateless RESTful API built on Express.js, communicating with a MongoDB Atlas cloud database via the Mongoose ODM.")
body("Key features of the current version include opportunity posting and browsing with advanced filters, a map-based spatial exploration view using Leaflet.js, a dynamic post lifecycle tracker, an AI-powered insights module, and a role-based access control system distinguishing between regular students and administrators. The project was developed iteratively following agile principles, with comprehensive API testing performed using Postman and multi-browser UI testing conducted to ensure cross-platform compatibility. MyCircle demonstrates a practical, scalable, and socially impactful application of modern web technologies in solving a real-world problem faced by the student community.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
h(1, "TABLE OF CONTENTS")
toc_items = [
    ("Certificate of Approval", "ii"),
    ("Guide Certificate", "iii"),
    ("Self Declaration", "iv"),
    ("Acknowledgement", "v"),
    ("Abstract", "vi"),
    ("Table of Contents", "vii"),
    ("List of Figures", "viii"),
    ("List of Abbreviations", "ix"),
    ("CHAPTER 1: Introduction", "1"),
    ("  1.1  Introduction to the Project", "1"),
    ("  1.2  Problem Statement", "3"),
    ("  1.3  Motivation", "4"),
    ("  1.4  Project Scope", "5"),
    ("  1.5  Organization of the Report", "6"),
    ("CHAPTER 2: Literature Review / Existing Systems", "7"),
    ("  2.1  Overview of Existing Platforms", "7"),
    ("  2.2  Comparative Analysis", "9"),
    ("  2.3  Research Gap", "10"),
    ("  2.4  Justification for Building MyCircle", "11"),
    ("CHAPTER 3: System Requirements", "12"),
    ("  3.1  Functional Requirements", "12"),
    ("  3.2  Non-Functional Requirements", "14"),
    ("  3.3  Hardware Requirements", "15"),
    ("  3.4  Software Requirements", "16"),
    ("CHAPTER 4: System Design", "17"),
    ("  4.1  System Architecture", "17"),
    ("  4.2  Architecture Diagram Description", "18"),
    ("  4.3  Database Design", "19"),
    ("  4.4  ER Diagram", "22"),
    ("  4.5  Data Flow Diagram (DFD)", "23"),
    ("  4.6  Use Case Diagram", "25"),
    ("  4.7  Module Description", "26"),
    ("  4.8  API Design", "28"),
    ("CHAPTER 5: Technology Stack", "32"),
    ("  5.1  MongoDB", "32"),
    ("  5.2  Express.js", "33"),
    ("  5.3  React.js", "34"),
    ("  5.4  Node.js", "36"),
    ("  5.5  Vite", "37"),
    ("  5.6  Tailwind CSS", "38"),
    ("  5.7  JWT Authentication", "39"),
    ("  5.8  Axios", "40"),
    ("  5.9  Socket.io", "41"),
    ("  5.10 Cloudinary", "42"),
    ("  5.11 Git & GitHub", "43"),
    ("CHAPTER 6: Implementation & Feature Walkthrough", "44"),
    ("  6.1  User Registration and Authentication", "44"),
    ("  6.2  Student Profile Page", "46"),
    ("  6.3  Opportunities Feed", "47"),
    ("  6.4  Post an Opportunity", "49"),
    ("  6.5  Campus Events", "50"),
    ("  6.6  Study Groups", "51"),
    ("  6.7  Follow / Unfollow System", "52"),
    ("  6.8  Notifications", "53"),
    ("  6.9  Messaging", "54"),
    ("  6.10 Admin Panel", "55"),
    ("CHAPTER 7: Testing", "57"),
    ("  7.1  Testing Strategy", "57"),
    ("  7.2  Unit Testing", "58"),
    ("  7.3  API Testing", "59"),
    ("  7.4  Frontend Testing", "62"),
    ("  7.5  Security Testing", "63"),
    ("  7.6  Performance Testing", "64"),
    ("  7.7  User Acceptance Testing", "65"),
    ("  7.8  Bug Report Table", "66"),
    ("CHAPTER 8: Screenshots / UI Walkthrough", "68"),
    ("CHAPTER 9: Limitations", "74"),
    ("CHAPTER 10: Future Scope", "77"),
    ("CHAPTER 11: Conclusion", "82"),
    ("References", "84"),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    from docx.oxml.ns import qn as _qn
    run = p.add_run(f"{title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(f"\t{page}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ══════════════════════════════════════════════════════════════════════════════
h(1, "LIST OF FIGURES")
figures = [
    ("Figure 4.1", "System Architecture Diagram – 3-Tier Client-Server Model", "18"),
    ("Figure 4.2", "Database Schema Overview – MongoDB Collections Relationship", "19"),
    ("Figure 4.3", "Entity-Relationship (ER) Diagram – MyCircle Database", "22"),
    ("Figure 4.4", "Level 0 DFD – Context Diagram (MyCircle System)", "23"),
    ("Figure 4.5", "Level 1 DFD – Decomposed Process View", "24"),
    ("Figure 4.6", "Use Case Diagram – Guest, Student, and Admin Actors", "25"),
    ("Figure 6.1", "Login and Registration Page UI", "44"),
    ("Figure 6.2", "Student Profile Page – View and Edit Modes", "46"),
    ("Figure 6.3", "Opportunities Feed – List View with Filters", "47"),
    ("Figure 6.4", "Map-Based Explore View – Leaflet Integration", "48"),
    ("Figure 6.5", "Post an Opportunity – Multi-step Form UI", "49"),
    ("Figure 6.6", "Campus Events Listing Page", "50"),
    ("Figure 6.7", "Study Groups – Browse and Create Interface", "51"),
    ("Figure 6.8", "Notification Bell and Dropdown Panel", "53"),
    ("Figure 6.9", "Direct Messaging Interface – Chat Window", "54"),
    ("Figure 6.10", "Admin Dashboard – User and Post Management", "55"),
    ("Figure 8.1", "Screenshot Placeholder – Login/Register Page", "68"),
    ("Figure 8.2", "Screenshot Placeholder – Home / Opportunity Feed", "69"),
    ("Figure 8.3", "Screenshot Placeholder – Profile Page", "70"),
    ("Figure 8.4", "Screenshot Placeholder – Post Opportunity Form", "71"),
    ("Figure 8.5", "Screenshot Placeholder – Events Page", "72"),
    ("Figure 8.6", "Screenshot Placeholder – Messaging Interface", "73"),
]
for fig, title, page in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{fig}  –  {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(f"\t{page}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ══════════════════════════════════════════════════════════════════════════════
h(1, "LIST OF ABBREVIATIONS")
abbrevs = [
    ("API", "Application Programming Interface"),
    ("MERN", "MongoDB, Express.js, React.js, Node.js"),
    ("JWT", "JSON Web Token"),
    ("REST", "Representational State Transfer"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("UI", "User Interface"),
    ("UX", "User Experience"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("HTTPS", "HyperText Transfer Protocol Secure"),
    ("DB", "Database"),
    ("JSON", "JavaScript Object Notation"),
    ("CORS", "Cross-Origin Resource Sharing"),
    ("SPA", "Single Page Application"),
    ("CDN", "Content Delivery Network"),
    ("MVC", "Model-View-Controller"),
    ("RBAC", "Role-Based Access Control"),
    ("ODM", "Object Document Mapper"),
    ("ORM", "Object Relational Mapper"),
    ("HMR", "Hot Module Replacement"),
    ("ESM", "ECMAScript Modules"),
    ("FCM", "Firebase Cloud Messaging"),
    ("UAT", "User Acceptance Testing"),
    ("DFD", "Data Flow Diagram"),
    ("ER", "Entity Relationship"),
    ("MCA", "Master of Computer Applications"),
    ("GGU", "Guru Ghasidas Vishwavidyalaya"),
    ("XSS", "Cross-Site Scripting"),
    ("2FA", "Two-Factor Authentication"),
    ("ML", "Machine Learning"),
    ("URL", "Uniform Resource Locator"),
]
add_table(["Abbreviation", "Full Form"], abbrevs, col_widths=[1.5, 4.5])

doc.add_page_break()

print("Part 1 complete — front matter done.")
doc.save("MyCircle_Report_TEMP.docx")
print("Saved temporary checkpoint.")
