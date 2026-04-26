
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document("MyCircle_Report_TEMP.docx")

def h(level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(max(12, 18 - level * 2))
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 1: INTRODUCTION")

h(2, "1.1 Introduction to the Project")
body("MyCircle is a community-driven, full-stack web application built on the MERN technology stack—MongoDB, Express.js, React.js, and Node.js—crafted specifically to serve the unmet networking and opportunity-discovery needs of students in local academic ecosystems. In an era when digital transformation has reshaped how professionals connect globally, students in tier-2 and tier-3 Indian cities still depend on informal WhatsApp groups, bulletin boards, and word-of-mouth to learn about internships, freelance projects, part-time jobs, hackathons, workshops, and study collaborations happening within their own campus and city. MyCircle is a deliberate, structured answer to this problem: a dedicated, verified, and hyper-local platform designed by a student, for students.")
body("The name 'MyCircle' is intentional and symbolic—it represents the idea that every student already belongs to multiple overlapping circles: their college circle, their city circle, their subject-interest circle, and their peer-follow circle. The platform brings all of these together under a single, unified digital space. Unlike global platforms that aggregate opportunities from across the country or world (and thereby drown out local, relevant posts), MyCircle surfaces only what matters to the student based on their institution, city, and interests. This locality-first philosophy is the platform's defining differentiator.")
body("The application supports two primary user roles: the Student and the Administrator. Students can register with institutional details, build a profile, post opportunities (internships, gigs, items for sale or rent), discover opportunities posted by peers, follow other students, communicate via a real-time direct messaging system, and receive event-driven notifications. The Admin role provides a governance layer with access to a management dashboard for user oversight, post moderation, and report handling. The frontend is a React.js Single Page Application (SPA) served by Vite, featuring a responsive, glassmorphism-inspired design system with smooth micro-animations powered by Framer Motion. The backend is a stateless RESTful API server built on Express.js, persisting all data in MongoDB Atlas via Mongoose, with JWT-based authentication securing all protected endpoints.")
body("MyCircle is not merely an academic exercise—it represents a genuinely deployable product. Real students can sign up, post opportunities, and connect with peers in their city today. The system has been tested, deployed on a cloud platform, and iterated upon based on real user feedback. It demonstrates the practical application of modern full-stack web development principles to a socially meaningful problem, and serves as a comprehensive showcase of the skills acquired during the MCA programme at Guru Ghasidas Vishwavidyalaya.")

h(2, "1.2 Problem Statement")
body("The core problem that MyCircle addresses can be stated as follows: students in local colleges and universities across India lack a centralized, verified, and community-specific digital platform to discover and post local academic and professional opportunities, connect with institutional peers, and collaborate on shared interests. This problem manifests in several interconnected ways that collectively impede student career development, academic collaboration, and campus community-building.")
body("First, the most widely used professional networking platform, LinkedIn, caters primarily to working professionals and job-seekers in the formal, large-company sector. Its algorithm and user base are oriented towards multinational corporations, metro-city recruiters, and white-collar roles. A small startup in Bilaspur looking for a part-time graphic designer intern from the local MCA batch has no effective means to target the right student audience on LinkedIn. Conversely, the MCA student looking for a local internship is lost in a sea of pan-India listings that are geographically and contextually irrelevant. The platform is not designed for hyperlocal, informal-sector, or college-circle networking.")
body("Second, existing opportunity-listing platforms such as Internshala and LinkedIn aggregators serve a discovery-only function—they list opportunities but provide no community infrastructure, no peer-to-peer following, no event management, no study-group formation, and no real-time communication. They are job boards, not networks. A student who finds an internship on Internshala cannot also see what events their college is hosting next weekend, join a Python study group, or message a senior alumnus in the same city for advice.")
body("Third, the informal workaround that most student communities resort to—WhatsApp groups, Telegram channels, Facebook groups—suffers from fundamental structural problems: there is no content categorization, no search functionality, no user verification, no post lifecycle management, and no discoverability for students outside the group. A student who recently joined the college cannot discover the 47 unrelated WhatsApp groups they need to join to stay informed about campus life. There is no single source of truth. MyCircle was designed to be exactly that single source of truth—structured, searchable, verified, and community-owned.")

h(2, "1.3 Motivation")
body("The motivation for building MyCircle emerged from a direct, personal observation of the opportunity-gap problem experienced by students at Guru Ghasidas Vishwavidyalaya. During the course of the MCA programme, it became evident that numerous valuable local opportunities—short-term freelance projects from local businesses, inter-department study collaborations, tech event announcements from nearby engineering colleges, and informal mentoring from senior students—were either never discovered by the students who needed them most, or were communicated so informally that they reached only a small subset of the intended audience. This recurring friction between opportunity availability and opportunity visibility was the primary motivational force behind MyCircle.")
body("From a technical standpoint, the motivation was equally compelling. The MCA curriculum at GGU provides a rigorous grounding in computer science fundamentals, data structures, database management, software engineering, and programming languages. However, the practical synthesis of these skills into a production-ready, full-stack web application represents a qualitatively different and higher-order learning challenge—one that involves not just writing code, but making architectural decisions, managing a RESTful API, implementing security mechanisms, designing a database schema, deploying to the cloud, and iterating based on user feedback. Building MyCircle was thus not only a solution to a real problem, but also a deliberate, ambitious attempt to consolidate and demonstrate the full breadth of technical skills acquired during the programme.")
body("There is also a broader motivational dimension rooted in the Indian higher education ecosystem. With over 900 universities and more than 40,000 colleges across India, the overwhelming majority of students study in institutions that lack robust career services, active alumni networks, and campus opportunity platforms. MyCircle represents a scalable template for how technology can democratize access to opportunities for students in these underserved academic ecosystems—not just in Bilaspur, but in every tier-2 and tier-3 city across the country.")

h(2, "1.4 Project Scope")
body("The current version of MyCircle, as documented in this report, encompasses a well-defined and deliberately bounded set of features designed to deliver core value while remaining achievable within the timeframe and resources of a Minor Project. The following features are included in scope for this version: user registration and JWT-based login; student profile creation and editing with profile image upload via Cloudinary; opportunity posting and browsing across four categories (jobs, services, sell, and rent); map-based spatial exploration using Leaflet.js integrated with the browser Geolocation API; a post lifecycle tracker with expiry management; the follow/unfollow peer-connection system; real-time one-on-one direct messaging; an in-app notification system; and a role-based admin panel for platform governance.")
body("The following features are explicitly out of scope for this version and are documented separately in the Future Scope chapter: a dedicated mobile application for iOS or Android (React Native is identified as the future framework); AI/ML-based personalized recommendation engines; multi-college institutional verification via institutional email domains (the current version uses self-declared college fields); a payment gateway for premium opportunity listings; a formal study group management module with group chat; email verification with OTP upon registration; and full Elasticsearch integration for advanced full-text search capabilities. These exclusions are not oversights but deliberate decisions to maintain project focus and ensure the in-scope features are implemented to a production-ready standard.")

h(2, "1.5 Organization of the Report")
body("The remainder of this report is organized into ten chapters, each addressing a distinct aspect of the MyCircle project. Chapter 2 presents a Literature Review, analysing existing platforms and identifying the research gap that MyCircle fills. Chapter 3 details the System Requirements, covering both functional and non-functional requirements alongside hardware and software specifications. Chapter 4 presents the System Design, encompassing the architectural overview, database schema, ER Diagram, Data Flow Diagrams (DFD), Use Case Diagram, module descriptions, and the complete API endpoint catalogue. Chapter 5 provides an in-depth Technology Stack analysis, explaining the rationale for each technology selected and how they integrate within MyCircle's architecture. Chapter 6 offers a feature-by-feature Implementation Walkthrough, describing the UI, underlying logic, and API interactions for each major module. Chapter 7 documents the Testing strategy and results, including API test cases, frontend testing, security testing, performance benchmarks, User Acceptance Testing (UAT), and a bug report. Chapter 8 presents the UI Walkthrough with screenshot descriptions for each major page. Chapter 9 identifies the current Limitations of the platform, and Chapter 10 outlines the Future Scope for enhancement and expansion. Chapter 11 concludes the report with a summary, learning outcomes, and closing statement. The References section lists all academic, technical, and online sources consulted during the project.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 2: LITERATURE REVIEW / EXISTING SYSTEMS")

h(2, "2.1 Overview of Existing Platforms")
body("A thorough examination of existing platforms that partially address student networking and opportunity discovery is essential to establish the context and necessity of MyCircle. This section evaluates four major categories of existing solutions: professional networking platforms (LinkedIn), opportunity-specific portals (Internshala), informal communication channels (WhatsApp groups), and college-specific institutional portals. Each category is analysed for its strengths, weaknesses, and specifically its failure to serve the hyperlocal, community-centric needs of students in smaller Indian cities.")
body("LinkedIn, founded in 2003 and currently the world's largest professional network with over 900 million members globally, is the default recommendation for professional networking. However, LinkedIn's design philosophy, algorithm, and content ecosystem are fundamentally oriented towards formal, corporate, and professional contexts. The platform privileges users with established work histories, endorsements from industry professionals, and connections to corporate entities. For a second-year MCA student in Bilaspur seeking a local part-time gig or a peer study partner for a Database Management Systems examination, LinkedIn offers little practical utility. Search results are dominated by listings from major metropolitan areas, and the social graph is sparse among students from non-premier institutions. Furthermore, LinkedIn's premium features—which include InMail messaging to non-connections and advanced job search filters—are behind a subscription paywall that is financially prohibitive for most Indian students. The platform's community features (LinkedIn Groups) have largely stagnated, and there is no mechanism for hyperlocal, institution-specific content surfacing.")
body("Internshala is India's leading internship-listing platform, with a user base primarily composed of undergraduate and postgraduate students. It provides a structured marketplace for internship and part-time job listings, with filters for location, stipend, and duration. While Internshala effectively addresses the opportunity-discovery dimension of the problem, it is fundamentally a job board, not a community platform. It offers no peer-to-peer follow system, no event management capability, no study group formation, and no real-time messaging between students. The listings are predominantly from companies seeking remote or metro-city candidates, and hyperlocal, informal opportunities from small local businesses that might be willing to hire a student from a nearby college are systematically absent. Internshala is a destination for formal internship applications, not a space for organic community interaction, collaborative learning, or peer-to-peer exchange.")
body("WhatsApp groups represent the most widely used informal solution to campus networking in Indian colleges. Virtually every college has dozens of WhatsApp groups—department groups, batch groups, event announcement groups, study material sharing groups, and alumni groups. These groups are omnipresent but deeply inefficient as an information infrastructure. There is no content categorization, making it impossible to search for a specific opportunity posted three weeks ago. There is no discovery mechanism for students to find relevant groups outside their immediate social circle. Content is volatile, with messages quickly buried under subsequent conversation. There is no user verification, meaning that anyone added to a group is indistinguishable from a genuine student. Group management is burdensome, with admins manually approving members and moderating content. WhatsApp groups are communication tools, not opportunity platforms, and they function as a poor substitute for a purpose-built student community application.")
body("Many universities and colleges maintain dedicated institutional portals or ERP systems for academic administration. However, these portals are overwhelmingly oriented towards administrative functions—fee payment, attendance records, examination results, and timetable management. Their social and community features, if any, are rudimentary and poorly adopted by students. The UX design of most institutional portals reflects their administrative, rather than community, purpose: they are functional but unappealing, with no social feed, no follow mechanism, no real-time features, and no mobile responsiveness. Furthermore, they are siloed—a student at GGU's engineering college cannot discover opportunities posted by a student at GGU's commerce department through the institutional portal.")

h(2, "2.2 Comparative Analysis")
body("The following table presents a structured comparative analysis of MyCircle against the four categories of existing platforms, evaluated on nine criteria relevant to the needs of local college students.")
add_table(
    ["Criteria", "LinkedIn", "Internshala", "WhatsApp Groups", "College Portals", "MyCircle"],
    [
        ["Target Audience", "Working Professionals", "Students (formal internships)", "Any group member", "Enrolled students only", "Local college students"],
        ["Hyperlocal Focus", "No", "Partial", "No (group-dependent)", "No", "Yes (city/college-level)"],
        ["Opportunity Posting", "Yes (formal jobs)", "Yes (internships/jobs)", "Informal", "No", "Yes (multi-category)"],
        ["Peer Follow System", "Yes (Connect)", "No", "No", "No", "Yes"],
        ["Real-time Messaging", "Yes (InMail, paid)", "No", "Yes", "No", "Yes (free)"],
        ["Event Management", "No", "No", "No", "Limited", "Yes"],
        ["Verified Student Profiles", "No", "No", "No", "Yes (admin)", "Yes (self-declared)"],
        ["Map-Based Discovery", "No", "No", "No", "No", "Yes (Leaflet.js)"],
        ["Cost to Student", "Free/Paid", "Free", "Free", "Free", "Free"],
    ],
    col_widths=[1.5, 1.1, 1.1, 1.2, 1.1, 1.1]
)

h(2, "2.3 Research Gap")
body("The comparative analysis reveals a clear and significant research gap: no existing platform simultaneously offers hyperlocal geographic focus, multi-category opportunity posting, a peer-follow social graph, real-time messaging, event management, and map-based spatial discovery—all within a single, freely accessible, student-verified platform. LinkedIn comes closest to providing a comprehensive social and professional network, but its hyperlocal and student-specific relevance is minimal. Internshala addresses opportunity discovery but offers no community infrastructure. WhatsApp groups enable communication but are structurally unfit for content organization and discoverability. College portals serve administrative functions but not community-building.")
body("The research gap can be precisely articulated as the absence of a platform that treats hyperlocal, institutional-community-specific student networking as a first-class problem, and builds every feature—from the social graph to the search filters to the map view—around the principle that what matters most to a student is what is happening within their academic circle, in their city, right now. MyCircle is explicitly designed to fill this gap.")

h(2, "2.4 Justification for Building MyCircle")
body("The justification for building MyCircle, rather than adapting or extending an existing platform, is grounded in the architectural and philosophical incompatibility between the needs of local student communities and the designs of existing solutions. Adapting LinkedIn would require overriding its globally-oriented algorithm, its professional-context bias, and its premium-feature paywall—a set of changes that would fundamentally alter the product and are not feasible for a student project. Building a plugin or extension for Internshala is not possible given its closed-source architecture. Creating a WhatsApp-based bot or structured group would replicate the communication capability but solve none of the discoverability, categorization, or social-graph problems.")
body("Building MyCircle from scratch using the MERN stack is therefore the most technically sound and practically impactful approach. It allows complete control over the data model, the API design, the user experience, and the feature set. It ensures that every design decision—from the choice of MongoDB's flexible document model to the implementation of JWT-based stateless authentication—is motivated by the specific requirements of a hyperlocal student community platform rather than by legacy constraints or commercial considerations. Furthermore, the act of building MyCircle from the ground up provides the developer with a comprehensive, hands-on learning experience spanning database design, API development, frontend engineering, authentication, cloud deployment, and real-time communication—precisely the breadth of practical skills that a full-stack MCA project is intended to cultivate.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 3: SYSTEM REQUIREMENTS")

h(2, "3.1 Functional Requirements")
body("Functional requirements define the specific behaviours and capabilities that the MyCircle system must provide to its users. The following functional requirements have been identified through a combination of stakeholder analysis (student community needs), competitive benchmarking (as detailed in Chapter 2), and iterative design exploration conducted during the project's planning phase.")
body("FR-01: User Registration — The system shall allow new users to register by providing their full name, email address, password, college name, city, and optionally a brief biography and profile photograph. The system shall validate that the email address is unique within the system before completing registration. Passwords shall be hashed using bcrypt before storage.")
body("FR-02: User Authentication — The system shall authenticate registered users via email and password. Upon successful authentication, the system shall issue a JSON Web Token (JWT) containing the user's ID and role, which the client shall store and include in the Authorization header of all subsequent API requests to protected endpoints.")
body("FR-03: Profile Management — Authenticated students shall be able to view and edit their profile information, including display name, biography, college, city, availability status, and profile photograph. Profile photographs shall be uploaded to Cloudinary, and the resulting URL shall be stored in the user document in MongoDB.")
body("FR-04: Opportunity Posting — Authenticated students shall be able to create opportunity posts specifying: title, post type (job, service, sell, rent), description, location (manual text or geolocation coordinates), price or budget range (minimum and maximum), duration, availability, tags, associated images (up to 5), whether barter is accepted, and an expiry date. The system shall persist these posts in the Opportunities collection.")
body("FR-05: Opportunity Discovery — The system shall provide an Explore feed displaying opportunity posts from across the platform, with support for filtering by post type, geographic proximity (using user-provided or GPS-detected location), search query (title and description text matching), and sorting by recency or distance.")
body("FR-06: Map-Based Exploration — The system shall render a map view (powered by Leaflet.js and OpenStreetMap tiles) displaying geolocated opportunity posts as interactive markers. Each marker shall display a category emoji and price badge, and clicking a marker shall reveal a popup with post summary and a link to the full post details page.")
body("FR-07: Follow/Unfollow System — Authenticated students shall be able to follow and unfollow other students. The follower and following counts shall be displayed on each user's profile page, along with paginated lists of followers and followees.")
body("FR-08: Real-time Direct Messaging — Authenticated students shall be able to initiate and participate in one-on-one direct message conversations with other students. Messages shall be persisted in the Messages collection. The messaging interface shall display a chronological conversation thread with read/delivered status indicators.")
body("FR-09: Notification System — The system shall generate and deliver in-app notifications for key user events, including new followers, likes and shares on the user's posts, new messages, and post expiry warnings. Notifications shall be marked as read when the user opens the notification panel.")
body("FR-10: Admin Panel — Users with the Administrator role shall have access to a dedicated admin dashboard displaying platform-wide statistics (total users, total posts, total active sessions). Administrators shall be able to view, deactivate, or delete user accounts, and remove or flag opportunity posts that violate community guidelines.")

h(2, "3.2 Non-Functional Requirements")
body("Non-functional requirements define the quality attributes and operational constraints of the MyCircle system. These requirements govern how the system performs its functions rather than what functions it performs.")
body("NFR-01: Performance — The system shall deliver initial page load times of under three seconds on a standard broadband connection (10 Mbps+). API endpoints for opportunity listing shall respond within 500 milliseconds for datasets of up to 10,000 posts. The React frontend, built with Vite and leveraging code-splitting and lazy loading, shall minimize the initial JavaScript bundle size to improve Time to Interactive (TTI).")
body("NFR-02: Security — All passwords shall be hashed using bcrypt with a minimum cost factor of 10. All protected API routes shall be guarded by JWT verification middleware. The system shall implement CORS policies restricting cross-origin requests to the designated frontend origin. Input sanitization shall be applied to all user-supplied data before it is processed by Mongoose queries, mitigating NoSQL injection risks. HTTPS shall be enforced on all production endpoints.")
body("NFR-03: Scalability — The MongoDB document model shall be designed to support horizontal scaling via MongoDB Atlas's auto-scaling capabilities. The Express.js API server shall remain stateless (all state stored in MongoDB or JWTs), enabling horizontal scaling across multiple server instances behind a load balancer. The system shall be designed to accommodate a user base scaling from hundreds to tens of thousands without architectural changes.")
body("NFR-04: Availability — The production deployment shall target 99%+ uptime, leveraging the reliability guarantees of the chosen cloud hosting platform (Render or Railway) and MongoDB Atlas. Graceful error handling shall ensure that individual component failures do not cascade into complete system unavailability.")
body("NFR-05: Usability — The frontend shall implement a mobile-first responsive design, ensuring full functionality and aesthetic integrity on screen widths from 320px (small mobile) to 1920px (large desktop). The UI shall conform to WCAG 2.1 Level AA accessibility guidelines where feasible, including appropriate colour contrast ratios, keyboard navigability, and semantic HTML structure.")
body("NFR-06: Maintainability — The codebase shall follow a modular architecture with clear separation of concerns: controllers, routes, models, middleware, and utility functions on the backend; pages, components, contexts, hooks, services, and utilities on the frontend. All API endpoints shall be documented in this report. Environment-specific configuration shall be managed via .env files and the platform's environment variable system.")

h(2, "3.3 Hardware Requirements")
body("The following table specifies the hardware requirements for both the development environment (the machine used to build and test MyCircle) and the production deployment environment (the cloud server on which MyCircle is hosted).")
add_table(
    ["Component", "Development Machine (Minimum)", "Production Server (Cloud)"],
    [
        ["Processor", "Intel Core i5 (8th Gen) / AMD Ryzen 5 or better", "2 vCPU (cloud instance, e.g., Render free tier)"],
        ["RAM", "8 GB DDR4 minimum (16 GB recommended)", "512 MB – 2 GB (based on plan)"],
        ["Storage", "256 GB SSD (for project files, Node modules)", "Ephemeral cloud storage (stateless server)"],
        ["Network", "Broadband Internet (10 Mbps+)", "Cloud datacenter network (Gbps)"],
        ["Display", "Full HD (1920×1080) recommended", "N/A (headless server)"],
        ["Operating System", "Windows 10/11, macOS 12+, or Ubuntu 20.04+", "Linux (Ubuntu 20.04 LTS, managed by cloud provider)"],
        ["Database", "MongoDB Atlas (cloud-hosted, no local server required)", "MongoDB Atlas M0 Free Tier (512 MB storage)"],
        ["Media Storage", "Cloudinary (cloud-hosted, no local disk needed)", "Cloudinary Free Tier (25 GB storage, 25 GB bandwidth/month)"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

h(2, "3.4 Software Requirements")
body("The following table catalogues all software tools, libraries, and platforms used in the development and deployment of MyCircle, along with the version used and the purpose each serves within the project.")
add_table(
    ["Software / Library", "Version", "Purpose"],
    [
        ["Node.js", "v18.17 LTS", "JavaScript runtime for backend server"],
        ["Express.js", "v4.18", "Web framework for RESTful API"],
        ["MongoDB", "v6.0 (Atlas)", "NoSQL cloud database"],
        ["Mongoose", "v7.4", "ODM for MongoDB schema and query management"],
        ["React.js", "v18.2", "Frontend UI library (component-based)"],
        ["Vite", "v5.0", "Frontend build tool and dev server"],
        ["Tailwind CSS", "v3.4", "Utility-first CSS framework for responsive design"],
        ["Framer Motion", "v10.16", "Animation library for React components"],
        ["React Router", "v6.15", "Client-side routing for React SPA"],
        ["Axios", "v1.5", "HTTP client for API requests from React"],
        ["Socket.io", "v4.6", "Real-time bidirectional event-based communication"],
        ["jsonwebtoken", "v9.0", "JWT generation and verification"],
        ["bcryptjs", "v2.4", "Password hashing and comparison"],
        ["Cloudinary SDK", "v1.41", "Image upload and CDN URL generation"],
        ["Multer", "v1.4", "Multipart form-data (file upload) middleware"],
        ["Leaflet.js", "v1.9", "Interactive map rendering"],
        ["React Leaflet", "v4.2", "React bindings for Leaflet"],
        ["cors", "v2.8", "CORS middleware for Express"],
        ["dotenv", "v16.3", "Environment variable management"],
        ["Postman", "v10.x", "API endpoint testing"],
        ["VS Code", "v1.85+", "Primary IDE for frontend and backend development"],
        ["Git", "v2.43", "Version control system"],
        ["GitHub", "N/A", "Remote repository and collaboration platform"],
        ["MongoDB Compass", "v1.41", "GUI for MongoDB database inspection and querying"],
        ["Google Chrome / Firefox", "Latest", "Browser for frontend testing and DevTools debugging"],
    ],
    col_widths=[2.0, 1.2, 3.3]
)

doc.add_page_break()

doc.save("MyCircle_Report_TEMP.docx")
print("Part 2 saved — Chapters 1, 2, 3 complete.")
