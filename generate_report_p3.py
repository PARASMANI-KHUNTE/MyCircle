
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document("MyCircle_Report_TEMP.docx")

def h(level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(max(12, 18 - level * 2))
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 4: SYSTEM DESIGN")

h(2, "4.1 System Architecture")
body("MyCircle is designed on a classic three-tier architecture, which separates the application into three distinct logical layers: the Presentation Tier (frontend client), the Logic Tier (backend API server), and the Data Tier (database). This separation of concerns ensures that each tier can be developed, tested, scaled, and maintained independently, providing a clean and modular system structure appropriate for a production-grade web application.")
body("The Presentation Tier is a React.js Single Page Application (SPA) bundled and served by Vite. When a user navigates to the MyCircle web application in their browser, the browser downloads the JavaScript, CSS, and HTML assets from the CDN or hosting service. The React application then takes control of the browser, rendering the UI, managing client-side routing via React Router v6, and communicating with the backend API exclusively via HTTP/HTTPS requests sent through the Axios library. The SPA architecture means that subsequent page navigations do not require full page reloads—React dynamically updates only the components that have changed, delivering a fast, app-like user experience.")
body("The Logic Tier is a stateless RESTful API server built on Node.js and Express.js. It receives HTTP requests from the React client, authenticates them via JWT middleware, processes the business logic (validating data, enforcing authorization rules, orchestrating data access), and returns JSON responses. Because the server is stateless—it stores no session information locally between requests—it is horizontally scalable: multiple server instances can run concurrently behind a load balancer, each independently processing requests. The Express.js application is organized into routes (defining endpoints), controllers (handling request/response logic), middleware (authentication, error handling, CORS), and utility modules (Cloudinary upload, token generation).")
body("The Data Tier consists of two cloud-hosted services: MongoDB Atlas for structured document storage, and Cloudinary for binary media asset storage. MongoDB Atlas is a fully managed cloud database service that hosts the MyCircle MongoDB database across distributed server nodes, providing automatic backups, high availability, and seamless scaling. The Express.js server connects to MongoDB Atlas via Mongoose, an ODM that provides schema definition, validation, query building, and middleware hooks. Cloudinary handles all user-uploaded image assets—profile photos and post images—storing them on its globally distributed CDN and returning publicly accessible URLs that are then stored in the MongoDB documents.")

h(2, "4.2 Architecture Diagram Description")
body("The system architecture of MyCircle, as represented in Figure 4.1, flows as follows. The User's Web Browser on the left side of the diagram represents the client layer. The browser loads the React SPA (JavaScript bundle built by Vite) and renders the UI. User interactions trigger Axios HTTP requests, which flow rightward over HTTPS to the Express.js API Server.")
body("The Express.js API Server sits at the centre of the architecture. Incoming requests first pass through the CORS Middleware, which validates the request origin. Authenticated requests then pass through the JWT Authentication Middleware, which decodes and verifies the Authorization header token. Validated requests are routed to the appropriate Controller, which implements the business logic for the requested operation. The controller interacts with Mongoose Models to read or write data in MongoDB Atlas (data flowing to the right). For media uploads, the controller uses the Cloudinary SDK to upload binary data to Cloudinary's servers and receive a CDN URL. For real-time events, the Socket.io server (co-hosted on the Express.js process) emits events to connected clients.")
body("MongoDB Atlas (top right) stores all structured application data across six primary collections: Users, Posts, Messages, Notifications, Conversations, and Requests. Cloudinary (bottom right) stores all media assets and serves them via its global CDN. The Socket.io layer (depicted as a bidirectional arrow between the browser and the Express server) maintains persistent WebSocket connections for real-time messaging and notification delivery.")

h(2, "4.3 Database Design")
body("MyCircle uses MongoDB as its database management system, leveraging its document-oriented, schema-flexible NoSQL data model. Unlike relational databases (MySQL, PostgreSQL) which store data in fixed-schema tables with rows and columns, MongoDB stores data as BSON (Binary JSON) documents within collections. This model is particularly well-suited to MyCircle's requirements because the opportunity post document, for example, can contain nested objects (location coordinates), arrays (tags, images, likes), and optional fields (budgetMin, budgetMax, expiresAt) without requiring schema alterations or NULL placeholders for every row as would be needed in a relational system.")
body("Mongoose, the Object Document Mapper used to interface between the Node.js application and MongoDB, provides schema definition, type enforcement, default values, validation, and middleware hooks (pre/post save, pre/post query) that give the flexible MongoDB document model a structured, application-level contract. Each Mongoose schema corresponds to a MongoDB collection. The following subsections describe each collection's schema in detail.")

h(3, "4.3.1 Users Collection")
body("The Users collection stores the core identity and social graph information for every registered user on the platform. Each user document contains the following fields: _id (ObjectId, auto-generated primary key), displayName (String, required, the user's chosen public name), email (String, required, unique, indexed for fast authentication lookups), passwordHash (String, required, the bcrypt-hashed password), college (String, optional, the user's institution), city (String, optional, the user's city), bio (String, optional, a short personal description), profileImage (String, optional, Cloudinary CDN URL of the profile photo), role (String, enum ['user', 'admin'], default 'user'), followers (Array of ObjectId references to Users), following (Array of ObjectId references to Users), reputation (Object containing trustScore and averageRating sub-fields), isActive (Boolean, default true, set to false when an admin deactivates the account), and createdAt / updatedAt (Date, auto-managed by Mongoose timestamps).")

h(3, "4.3.2 Posts Collection")
body("The Posts collection is the central data collection of MyCircle, storing all opportunity posts created by students. Each post document contains: _id (ObjectId), title (String, required), description (String, required), type (String, enum ['job', 'service', 'sell', 'rent'], required), location (String, human-readable location name), coordinates (Object with lat and lng sub-fields for geospatial queries), price (Number, optional), budgetMin (Number, optional), budgetMax (Number, optional), images (Array of Strings, Cloudinary URLs), tags (Array of Strings), user (ObjectId reference to Users, the post author), availability (String, optional), duration (Number, days), acceptsBarter (Boolean, default false), isUrgent (Boolean, default false), isActive (Boolean, default true), status (String, enum ['active', 'expired', 'completed', 'cancelled']), expiresAt (Date, computed from createdAt + duration), likes (Array of ObjectId references to Users who liked the post), shares (Number, share count), views (Number, view count), and createdAt / updatedAt timestamps.")

h(3, "4.3.3 Messages Collection")
body("The Messages collection stores individual chat messages exchanged between users. Each message document contains: _id (ObjectId), conversation (ObjectId reference to Conversations), sender (ObjectId reference to Users), content (String, the message text), messageType (String, enum ['text', 'system', 'transaction'], default 'text'), isRead (Boolean, default false), createdAt (Date), and optionally transactionData (Object containing amount, description, and status sub-fields for in-chat transaction records).")

h(3, "4.3.4 Conversations Collection")
body("The Conversations collection acts as a container for messaging threads between pairs of users. Each document contains: _id (ObjectId), participants (Array of two ObjectId references to Users), post (ObjectId reference to Posts, the opportunity that initiated the conversation), lastMessage (ObjectId reference to the most recent Message), lastMessageTime (Date), unreadCount (Object mapping user IDs to their respective unread message counts), and createdAt / updatedAt timestamps.")

h(3, "4.3.5 Notifications Collection")
body("The Notifications collection stores in-app notifications for all user events. Each notification document contains: _id (ObjectId), recipient (ObjectId reference to Users), sender (ObjectId reference to Users, optional), type (String, enum ['follow', 'like', 'share', 'message', 'post_expiry', 'system']), message (String, the notification display text), relatedPost (ObjectId reference to Posts, optional), isRead (Boolean, default false), and createdAt timestamp.")

h(3, "4.3.6 Requests Collection")
body("The Requests collection manages opportunity-application and peer-connection requests between students. Each request document contains: _id (ObjectId), fromUser (ObjectId reference to Users), toUser (ObjectId reference to Users), post (ObjectId reference to Posts), status (String, enum ['pending', 'approved', 'rejected', 'cancelled']), message (String, the applicant's introductory message), and createdAt / updatedAt timestamps. When a request is approved by the post owner, a Conversation document is automatically created between the two parties, enabling direct messaging.")

h(2, "4.4 ER Diagram")
body("Figure 4.3 depicts the Entity-Relationship Diagram for the MyCircle database, illustrating the six primary entities and their relationships. The following description is intended to guide the manual drawing of the ER Diagram on plain paper or using a diagramming tool.")
body("The USER entity is the central entity of the diagram, with attributes: UserID (PK), DisplayName, Email, PasswordHash, College, City, Bio, ProfileImage, Role, IsActive, CreatedAt. The USER entity has the following relationships: a USER can create zero or many POSTS (one-to-many, 'Posts' relationship); a USER can send zero or many MESSAGES (one-to-many); a USER can receive zero or many MESSAGES (one-to-many); a USER can follow zero or many other USERs, and be followed by zero or many USERs (many-to-many self-referencing, 'Follows' relationship); a USER can receive zero or many NOTIFICATIONS (one-to-many); a USER can be the initiator or recipient of zero or many REQUESTS (one-to-many).")
body("The POST entity has attributes: PostID (PK), Title, Description, Type, Location, Coordinates, Price, BudgetMin, BudgetMax, Tags, Images, Duration, ExpiresAt, Status, IsActive, CreatedAt. Each POST is associated with exactly one USER (the author). A POST can have zero or many REQUESTS associated with it. A POST can be part of zero or many CONVERSATIONS (one-to-many).")
body("The MESSAGE entity has attributes: MessageID (PK), Content, MessageType, IsRead, CreatedAt. Each MESSAGE belongs to exactly one CONVERSATION, is sent by exactly one USER (sender). The CONVERSATION entity has attributes: ConversationID (PK), LastMessageTime, UnreadCount. Each CONVERSATION has exactly two participants (USERs) and is optionally associated with one POST. The NOTIFICATION entity has attributes: NotificationID (PK), Type, Message, IsRead, CreatedAt. Each NOTIFICATION has one recipient USER and optionally one sender USER and one related POST. The REQUEST entity has attributes: RequestID (PK), Status, Message, CreatedAt. Each REQUEST is from one USER (fromUser) to another USER (toUser) and is associated with one POST.")

h(2, "4.5 Data Flow Diagram (DFD)")

h(3, "4.5.1 Level 0 DFD (Context Diagram)")
body("The Level 0 DFD, also known as the Context Diagram (Figure 4.4), represents MyCircle as a single process — 'MyCircle Platform' — at the centre of the diagram. There are three external entities that interact with the system: Student User (the primary actor), Administrator, and External Services (Cloudinary for image hosting, MongoDB Atlas for cloud data persistence). The data flows are as follows: The Student User sends Registration Data, Login Credentials, Opportunity Post Data, Message Content, and Follow/Unfollow Requests to the MyCircle Platform. The MyCircle Platform returns Authentication Tokens, Opportunity Feed Data, Profile Data, Notification Data, and Chat Messages to the Student User. The Administrator sends Moderation Actions (delete post, deactivate user) and receives Platform Statistics and User/Post Lists. The MyCircle Platform sends Image Upload Requests to External Services (Cloudinary) and receives CDN Image URLs in return. The Platform sends Database Read/Write Queries to MongoDB Atlas and receives Document Data.")

h(3, "4.5.2 Level 1 DFD")
body("The Level 1 DFD (Figure 4.5) decomposes the MyCircle Platform process into seven sub-processes, each with their associated data flows and data stores. The seven processes are: P1 (User Authentication), P2 (Profile Management), P3 (Opportunity Management), P4 (Social Graph Management), P5 (Messaging), P6 (Notification System), and P7 (Admin Panel).")
body("P1 — User Authentication: receives Registration Data and Login Credentials from the Student User external entity. Writes new User documents to the DS1 (Users Data Store). Returns JWT Access Token to the client. P2 — Profile Management: receives Profile Update Data and Image Upload from the Student User. Reads from and writes to DS1 (Users Data Store). Sends image data to Cloudinary and receives CDN URL. P3 — Opportunity Management: receives Post Creation Data and Search/Filter Queries from the Student User. Reads from and writes to DS2 (Posts Data Store). Returns Opportunity Feed Data and Post Detail Data to the client. P4 — Social Graph Management: receives Follow/Unfollow Requests. Reads from and writes to DS1 (Users Data Store) to update follower/following arrays. Returns updated social graph data. P5 — Messaging: receives Message Content from Student User. Reads from and writes to DS3 (Messages Data Store) and DS4 (Conversations Data Store). Returns Conversation Thread Data to both the sender and recipient in real-time via Socket.io. P6 — Notification System: triggered by events from P3, P4, and P5. Writes notification documents to DS5 (Notifications Data Store). Delivers real-time notification events to the recipient client via Socket.io. P7 — Admin Panel: receives Moderation Actions from the Administrator. Reads from DS1 (Users), DS2 (Posts), DS5 (Notifications). Writes deactivation/deletion updates. Returns Platform Statistics to the Administrator.")

h(2, "4.6 Use Case Diagram")
body("Figure 4.6 depicts the Use Case Diagram for MyCircle. The diagram contains three actors: Guest (unauthenticated visitor), Student (authenticated registered user), and Admin (authenticated administrator). The use cases are organized within the system boundary (a rectangle labelled 'MyCircle System').")
body("Guest Actor Use Cases: (1) View Landing/Home Page, (2) Register New Account, (3) Log In with Credentials, (4) Browse Public Opportunity Feed (limited view). The Guest actor can only interact with the publicly accessible portions of the platform. All other operations require authentication.")
body("Student Actor Use Cases (extends Guest): (5) View and Edit Own Profile, (6) Upload Profile Photo, (7) Create Opportunity Post, (8) Edit Own Opportunity Post, (9) Delete Own Opportunity Post, (10) Browse Opportunity Feed with Filters, (11) View Map-Based Opportunity Exploration, (12) View Full Opportunity Details, (13) Like an Opportunity Post, (14) Share an Opportunity Post, (15) Send Application Request for an Opportunity, (16) Follow Another Student, (17) Unfollow a Student, (18) View Follower/Following Lists, (19) Send Direct Message, (20) View Conversation Thread, (21) Receive Real-time Notifications, (22) Mark Notifications as Read, (23) View AI Insights for a Post, (24) View Post Analytics (own posts only).")
body("Admin Actor Use Cases (extends Student): (25) View Admin Dashboard with Platform Statistics, (26) View All Users List, (27) Deactivate/Reactivate User Account, (28) Delete User Account, (29) View All Posts List, (30) Delete Any Post, (31) View Reported Content, (32) Manage Platform Settings. The Admin actor has all Student capabilities plus the additional governance capabilities listed above. The include relationship exists between 'Create Opportunity Post' and 'Upload Post Images' (image upload is an included sub-use-case of post creation). The extend relationship exists between 'Browse Opportunity Feed' and 'View Map Exploration' (map view is an optional extension of browsing).")

h(2, "4.7 Module Description")
body("MyCircle's architecture is organized into eight functional modules, each encapsulating a cohesive set of features and corresponding backend routes, frontend pages, and database interactions. The following table provides a structured overview of each module.")
add_table(
    ["Module", "Description", "Key Frontend Components", "Key Backend Routes"],
    [
        ["Auth Module", "Handles user registration, login, JWT issuance, and session management. Implements bcrypt hashing and JWT verification middleware.", "LoginPage, RegisterPage, AuthContext, ProtectedRoute", "/api/auth/register, /api/auth/login, /api/auth/me"],
        ["User/Profile Module", "Manages student profile display, editing, and social graph (followers/following). Handles profile image upload to Cloudinary.", "ProfilePage, EditProfileModal, FollowersList", "/api/users/:id, /api/users/:id/follow, /api/users/:id/followers"],
        ["Opportunities Module", "Core post management: creating, reading, updating, deleting opportunity posts. Includes the Explore feed with filters and map view.", "Explore, PostCard, PostDetails, CreatePost, EditPost", "/api/posts (CRUD), /api/posts/nearby, /api/posts/:id/like"],
        ["Requests Module", "Manages application requests from interested students to post owners. Triggers conversation creation on approval.", "Requests, RequestCard", "/api/requests, /api/requests/:id/approve, /api/requests/:id/reject"],
        ["Messaging Module", "Real-time one-on-one chat between students. Manages conversations and individual messages via Socket.io.", "Messages, ChatWindow, ConversationList", "/api/conversations, /api/conversations/:id/messages"],
        ["Notifications Module", "Generates, delivers, and manages in-app notifications for user events. Uses Socket.io for real-time delivery.", "NotificationBell, NotificationPanel", "/api/notifications, /api/notifications/read-all"],
        ["Admin Module", "Platform governance: user management, post moderation, statistics dashboard.", "AdminDashboard, UserManagement, PostModeration", "/api/admin/stats, /api/admin/users, /api/admin/posts"],
        ["AI Insights Module", "Provides AI-generated insights for post owners (market demand, price analysis) and viewers (post summary, context). Calls external AI API.", "AIInsightsPanel (in PostCard)", "/api/posts/:id/insights, /api/posts/:id/explanation"],
    ],
    col_widths=[1.2, 2.0, 1.8, 1.5]
)

h(2, "4.8 API Design")
body("MyCircle's backend implements a RESTful API, adhering to the principles of Representational State Transfer (REST). The API uses standard HTTP methods (GET, POST, PUT, PATCH, DELETE) to perform CRUD operations on resources. All requests and responses use the JSON content type. Protected endpoints require a valid JWT in the Authorization header (Bearer <token>). The base URL for all API endpoints is: https://mycircle-api.onrender.com/api (production) or http://localhost:5000/api (development).")
body("The following table documents all major API endpoints implemented in the current version of MyCircle. For brevity, request body schemas are described in prose rather than full JSON Schema format. 'Y' in the Auth column indicates that a valid JWT Bearer token is required.")
add_table(
    ["Method", "Endpoint", "Description", "Auth"],
    [
        ["POST", "/auth/register", "Register a new user account. Body: displayName, email, password, college, city.", "N"],
        ["POST", "/auth/login", "Authenticate user with email/password. Returns JWT and user object.", "N"],
        ["GET", "/auth/me", "Retrieve the authenticated user's profile using their JWT.", "Y"],
        ["GET", "/users/:id", "Retrieve a specific user's public profile by their MongoDB ObjectId.", "N"],
        ["PUT", "/users/:id", "Update the authenticated user's profile (name, bio, college, city, profileImage URL).", "Y"],
        ["POST", "/users/:id/follow", "Toggle follow/unfollow for the user identified by :id.", "Y"],
        ["GET", "/users/:id/followers", "Retrieve paginated list of users who follow the specified user.", "N"],
        ["GET", "/users/:id/following", "Retrieve paginated list of users the specified user follows.", "N"],
        ["GET", "/posts", "Retrieve paginated, filterable list of opportunity posts. Query params: type, city, search, page, limit, sortBy.", "N"],
        ["POST", "/posts", "Create a new opportunity post. Body: title, type, description, location, coordinates, price, budgetMin, budgetMax, images, tags, duration, availability, acceptsBarter, isUrgent.", "Y"],
        ["GET", "/posts/nearby", "Retrieve posts within a radius of given coordinates. Query: lat, lng, radius.", "N"],
        ["GET", "/posts/:id", "Retrieve full details of a specific post by ID. Increments view count.", "N"],
        ["PUT", "/posts/:id", "Update an existing post (author only). Same body schema as POST /posts.", "Y"],
        ["DELETE", "/posts/:id", "Delete a post (author or admin only). Sets isActive to false.", "Y"],
        ["POST", "/posts/:id/like", "Toggle like/unlike on a post for the authenticated user.", "Y"],
        ["POST", "/posts/:id/share", "Record a share action and increment the post's share count.", "Y"],
        ["GET", "/posts/:id/analytics", "Retrieve view, like, share, and days-active statistics for a post (owner only).", "Y"],
        ["GET", "/conversations", "Retrieve all conversations for the authenticated user.", "Y"],
        ["POST", "/conversations", "Create a new conversation (or retrieve existing) between two users for a post. Body: recipientId, postId.", "Y"],
        ["GET", "/conversations/:id/messages", "Retrieve paginated message history for a specific conversation.", "Y"],
        ["POST", "/conversations/:id/messages", "Send a new message in a conversation. Body: content, messageType.", "Y"],
        ["GET", "/notifications", "Retrieve all unread and recent notifications for the authenticated user.", "Y"],
        ["PATCH", "/notifications/read-all", "Mark all notifications for the authenticated user as read.", "Y"],
        ["GET", "/requests", "Retrieve all requests (sent or received) for the authenticated user.", "Y"],
        ["POST", "/requests", "Create a new application request for an opportunity. Body: postId, toUserId, message.", "Y"],
        ["PATCH", "/requests/:id/approve", "Approve a received request, triggering conversation creation (post owner only).", "Y"],
        ["PATCH", "/requests/:id/reject", "Reject a received request (post owner only).", "Y"],
        ["GET", "/admin/stats", "Retrieve platform-wide statistics: total users, posts, conversations. Admin only.", "Y"],
        ["GET", "/admin/users", "Retrieve paginated list of all users. Admin only.", "Y"],
        ["PATCH", "/admin/users/:id/deactivate", "Deactivate or reactivate a user account. Admin only.", "Y"],
        ["DELETE", "/admin/posts/:id", "Hard-delete any post. Admin only.", "Y"],
    ],
    col_widths=[0.6, 2.0, 2.6, 0.4]
)

doc.add_page_break()

doc.save("MyCircle_Report_TEMP.docx")
print("Part 3 saved — Chapter 4 complete.")
