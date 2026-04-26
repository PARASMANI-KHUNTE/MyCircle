from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document("MyCircle_Report_TEMP.docx")

def h(level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(max(12, 18 - level * 2))
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 — IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 6: IMPLEMENTATION & FEATURE WALKTHROUGH")

h(2, "6.1 User Registration and Authentication")
body("The authentication system is the foundational security layer of MyCircle. The registration flow begins with the Register page, which presents a clean, multi-field form collecting the user's Display Name, Email Address, Password, College Name, and City. Client-side validation is implemented using controlled React components: the email field validates format using a regular expression, the password field enforces a minimum length of 8 characters and provides a real-time strength indicator, and all required fields display inline error messages if submission is attempted while they are empty. The form uses a loading state to disable the submit button and show a spinner while the API call is in progress, preventing duplicate submissions.")
body("Upon form submission, the React component calls api.post('/auth/register', formData). On the server, the auth controller first checks whether a user with the submitted email already exists in the Users collection. If so, it returns a 409 Conflict response. If not, it hashes the password using bcryptjs (await bcrypt.hash(password, 10)) and creates a new User document in MongoDB. The server then immediately issues a JWT so the user is logged in upon successful registration, avoiding the friction of a separate login step. The JWT is stored in the React application's AuthContext state and in localStorage for persistence across browser sessions.")
body("The login flow mirrors the registration flow. The Login page presents a two-field form (email and password). Upon submission, the credentials are sent to POST /api/auth/login. The server retrieves the user by email, uses bcrypt.compare() to verify the password against the stored hash, and returns a signed JWT on success. Protected routes in React are implemented using a ProtectedRoute wrapper component that reads the authentication state from AuthContext and redirects unauthenticated users to the Login page. The Axios instance is configured with a response interceptor that catches 401 Unauthorized responses and automatically clears the stored token and redirects the user to login, handling token expiry gracefully without user confusion.")

h(2, "6.2 Student Profile Page")
body("The Profile page serves as the student's public identity card within the MyCircle community. It displays the user's profile photograph (served from Cloudinary CDN), display name, college, city, biography, account creation date, and reputation metrics (trust score and average rating). The social graph is prominently displayed: follower and following counts are clickable, opening a modal that lists the respective users with their avatars and names. Below the social metrics, the profile displays all opportunity posts authored by the user in a responsive card grid, allowing visitors to quickly assess the user's posting history and community engagement.")
body("The Edit Profile functionality is accessible only to the authenticated profile owner (determined by comparing the URL's userId parameter with the authenticated user's ID from AuthContext). Clicking the Edit button opens an inline edit form or modal that pre-populates all fields with the current profile data. The profile image upload field uses an HTML file input with an accept='image/*' attribute, and upon file selection, the component creates a FormData object containing the image file and the other profile fields, sending it to PUT /api/users/:id with the Content-Type: multipart/form-data header. The server processes the upload through Multer, uploads the image buffer to Cloudinary, and stores the returned CDN URL in the user document.")
body("The follow/unfollow button on another user's profile is implemented as a single toggle button whose visual state (outlined 'Follow' vs. filled 'Following') is derived from whether the authenticated user's ID is present in the viewed profile's followers array. Clicking it calls POST /api/users/:id/follow, which uses a MongoDB $addToSet / $pull operation to atomically add or remove the follower relationship, preventing race conditions. The profile page's follower count and button state update optimistically in the UI without requiring a full page reload.")

h(2, "6.3 Opportunities Feed (Explore Page)")
body("The Explore page is the most complex and feature-rich page in MyCircle. It is organized into two primary view modes: the List View (a responsive card grid of opportunity posts) and the Map View (an interactive Leaflet map with geolocated post markers). A view toggle at the top of the page switches between these two modes. The page also contains a comprehensive filter panel with controls for post type (Job, Service, Sell, Rent), search query (full-text match on title and description), and geographic proximity (a radius filter centered on the user's GPS location).")
body("In List View, opportunities are displayed as PostCard components arranged in a responsive CSS Grid (1 column on mobile, 2 columns on tablet, 3 columns on desktop). Each PostCard shows the post's cover image (or a category emoji placeholder), post type badge, title, author avatar and name, posting date, location, price/budget badge, and a footer with action buttons (AI Insights, Like, Share, Analytics for own posts, and a More/Less toggle for the collapsible details section). The card is designed to be compact by default, showing only the most essential information, with Lifecycle, Budget Tags, and Edit/Delete controls hidden until the user clicks the 'More' chevron button. This collapsible architecture, implemented using Framer Motion's AnimatePresence, keeps the feed scannable and reduces visual clutter.")
body("In Map View, the Leaflet map renders custom pin-shaped markers for each geolocated post. Each marker displays the category emoji and a price badge. Clicking a marker opens a popup with the post title, type badge, price, and a 'View Details' link. A glassmorphism-styled control panel overlaid on the map provides Zoom In/Out buttons, a Recenter button (which flies the map to the user's GPS location), and a Radius Toggle button (which shows/hides a semi-transparent circle indicating the search radius around the user). The map view is an excellent tool for students who want to discover what opportunities are physically close to them in their city.")

h(2, "6.4 Post an Opportunity")
body("The Create Post page presents a multi-section form for creating a new opportunity post. The form is logically organized into sections: Basic Information (Title, Post Type selector, Description), Pricing (Price, Budget Min/Max, Accepts Barter toggle), Details (Duration, Availability, Tags input with comma-separated parsing), Location (a text field for location name plus an optional interactive Leaflet map for pin-pointing coordinates), Images (a multi-file upload field that previews selected images as thumbnails), and Settings (Is Urgent toggle, Expiry Duration). The Post Type selector uses a visual card-style selector where each type (Job, Service, Sell, Rent) is displayed with its emoji and label, and selecting a type highlights the corresponding card with a primary colour border.")
body("Client-side validation ensures that Title, Post Type, and Description are provided before submission. The form uses a FormData object to package both the text fields and image files for transmission to POST /api/posts. On the server, the controller validates all required fields, uploads any provided images to Cloudinary (iterating through the array of file buffers), computes the expiresAt date from the duration value, and inserts the new post document into the Posts collection. Upon successful creation, the user is navigated to the new post's detail page, and a success toast notification confirms the action.")

h(2, "6.5 Campus Events")
body("While a full Events module with dedicated database collection is identified as future scope, the current version of MyCircle supports event-type opportunities through the existing Posts collection with the 'service' or 'job' type. Upcoming iterations will introduce a dedicated Events collection with fields for event date, venue, organizer, and registration link. In the current version, students post campus events as opportunity cards with the relevant details in the description field and use tags to categorize the content as an event. The map view and proximity filter remain applicable, allowing students to discover events happening near their location.")

h(2, "6.6 Study Groups")
body("Similar to Events, the Study Groups feature is identified as a future scope enhancement in the current version. Students currently leverage the Direct Messaging and the Requests system to form informal study collaborations. A dedicated Study Groups module—with its own collection, group creation form, member management, and group chat—is documented in Chapter 10 as a prioritized future enhancement. The architectural groundwork (Socket.io real-time layer, conversation infrastructure) is already in place to support group chat implementation.")

h(2, "6.7 Follow / Unfollow System")
body("The social graph in MyCircle is built on a straightforward follow/unfollow mechanism. When Student A follows Student B, A's user document gains B's ObjectId in its 'following' array, and B's user document gains A's ObjectId in its 'followers' array. This bidirectional update is performed atomically using MongoDB's $addToSet operator (to prevent duplicate entries) in a single findByIdAndUpdate call per user document. The unfollow action uses $pull to remove the ObjectId from both arrays. Following another student does not require the followed user's approval—it is a one-directional follow model (similar to Twitter/Instagram) rather than a mutual-connection model (like LinkedIn). This design choice enables open content discovery: a student can follow any other student's posts without requiring reciprocation, encouraging broader community engagement.")
body("The follow action also triggers a notification. When A follows B, the Notification controller creates a new document in the Notifications collection with type 'follow', recipient B, and sender A. The Socket.io server then emits a 'new_notification' event to B's socket room, causing B's notification bell to update in real-time without any polling.")

h(2, "6.8 Notifications")
body("The Notifications system in MyCircle provides in-app, real-time alerts for user events. The notification bell icon in the navigation bar displays a badge with the count of unread notifications. Clicking the bell opens a dropdown panel listing the most recent notifications, each with an icon corresponding to its type (heart for like, user-plus for follow, message-circle for message, sparkles for AI insight), a descriptive text message, the sender's avatar, and the time elapsed since the notification was created. Clicking a notification marks it as read and navigates to the relevant content (the post that was liked, the profile of the new follower, etc.).")
body("Notifications are generated on the server by the Notification service module, which is called by the relevant controllers (Posts controller for like/share events, Users controller for follow events, Conversations controller for message events). Each notification is persisted in MongoDB and simultaneously emitted to the recipient's Socket.io room. The frontend Socket context listens for 'new_notification' events and updates the notification count and list in real-time. The 'Mark All as Read' button calls PATCH /api/notifications/read-all, which uses a MongoDB updateMany() to set isRead: true on all notifications for the authenticated user.")

h(2, "6.9 Messaging")
body("The messaging system in MyCircle enables one-on-one direct communication between students. A conversation is typically initiated from a post's detail page, where an interested student clicks 'Send Message' or 'Request', which creates a Conversation document linking the two users and the related post. The Messages page displays a list of all the authenticated user's conversations, each showing the other participant's avatar, name, the last message preview, and the time of the last message. Unread conversations are visually distinguished with a bold last-message preview and a coloured dot indicator.")
body("The ChatWindow component renders the full conversation thread in a scrollable container, with messages from the authenticated user aligned to the right (in the primary colour bubble) and messages from the other participant aligned to the left (in a neutral card-coloured bubble). The message input at the bottom is a resizable textarea with a Send button. When the user sends a message, the component calls POST /api/conversations/:id/messages, which saves the message to MongoDB and emits a 'new_message' Socket.io event to the recipient's room. The recipient's ChatWindow (if open) receives the event and appends the new message to the thread in real-time. Read receipts are implemented: when a user opens a conversation, all messages from the other participant are marked as read via a PATCH request.")

h(2, "6.10 Admin Panel")
body("The Admin Panel is accessible only to users with the 'admin' role, enforced by both the frontend ProtectedRoute (which checks the role from AuthContext) and the backend admin middleware (which verifies the role from the JWT payload). The Admin Dashboard presents a statistics overview: total registered users, total active posts, total conversations, and a trend indicator showing growth since the previous period. Below the statistics, the dashboard provides two management tabs: User Management and Post Management.")
body("The User Management tab displays a paginated table of all registered users, with columns for avatar, name, email, college, role, registration date, and account status (active/deactivated). Each row has action buttons for Deactivate/Reactivate (which toggles the user's isActive flag, preventing login without deleting the account) and Delete (which permanently removes the user document). The Post Management tab displays a paginated table of all posts with columns for title, type, author, creation date, status, and view/like/share counts. Administrators can delete any post directly from this table. These moderation capabilities ensure that the platform remains free of spam, inappropriate content, and policy-violating posts, maintaining a safe and constructive community environment.")

doc.add_page_break()
doc.save("MyCircle_Report_TEMP.docx")
print("Part 5 saved — Chapter 6 complete.")
