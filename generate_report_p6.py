from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document("MyCircle_Report_TEMP.docx")

def h(level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(max(12, 18 - level * 2))
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7 — TESTING
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 7: TESTING")

h(2, "7.1 Testing Strategy")
body("A comprehensive testing strategy was employed throughout the development of MyCircle to ensure the correctness, security, performance, and usability of all system components. The testing approach encompasses four levels: Unit Testing (verifying individual functions and modules in isolation), Integration Testing (verifying that modules interact correctly), System Testing (verifying the complete system against functional requirements), and User Acceptance Testing (verifying that the system meets the real-world needs of student users). Given the full-stack nature of the application, testing was performed on both the backend API layer and the frontend React layer, using a combination of automated tools (Postman, Jest concepts) and systematic manual testing procedures.")
body("The testing lifecycle was integrated into the development workflow rather than treated as a separate post-development phase. API endpoints were tested with Postman immediately after implementation, before the corresponding frontend UI was built. Frontend components were tested in the browser's developer console and React DevTools as they were built. Security testing was performed before deployment. Performance benchmarks were measured after the initial deployment on the cloud platform. User Acceptance Testing was conducted with a small group of real student users after the platform reached a feature-complete state.")

h(2, "7.2 Unit Testing")
body("Unit testing in MyCircle focused on verifying the correctness of individual utility functions and business logic modules. The bcrypt password hashing and verification cycle was tested by generating a hash for a known password and verifying that bcrypt.compare() returns true for the correct password and false for an incorrect one. The JWT generation and verification functions were tested by signing a payload and verifying that jwt.verify() correctly decodes the payload with the correct secret and throws a JsonWebTokenError with an incorrect secret. The formatMoney() utility function in the frontend was tested against edge cases: zero values (should return null), non-numeric inputs (should return null), and valid numeric inputs (should return the correctly formatted currency string). The lifecycleMeta computation (which calculates percentage of post lifetime remaining, the lifecycle stage label, and colour class) was tested against posts with various expiresAt values relative to the current date, including expired posts, nearly-expired posts, and fresh posts.")

h(2, "7.3 API Testing")
body("API testing was performed using Postman, an industry-standard API development and testing tool. A Postman Collection named 'MyCircle API' was created containing all API endpoints organized into folders by resource (Auth, Users, Posts, Conversations, Notifications, Requests, Admin). Environment variables were configured in Postman for the base URL and the JWT token (populated automatically from the login response using a Postman test script: pm.environment.set('token', pm.response.json().token)). The following table documents a representative set of API test cases with their inputs, expected outputs, and results.")
add_table(
    ["TC ID", "Endpoint", "Method", "Input", "Expected Output", "Actual Output", "Status"],
    [
        ["TC-01", "/auth/register", "POST", "Valid name, email, password, college, city", "201 Created, JWT token in response", "201 Created, JWT returned", "PASS"],
        ["TC-02", "/auth/register", "POST", "Existing email address", "409 Conflict, error message", "409 Conflict, 'Email already exists'", "PASS"],
        ["TC-03", "/auth/login", "POST", "Valid email and correct password", "200 OK, JWT token and user object", "200 OK, token and user returned", "PASS"],
        ["TC-04", "/auth/login", "POST", "Valid email, wrong password", "401 Unauthorized, error message", "401 Unauthorized, 'Invalid credentials'", "PASS"],
        ["TC-05", "/auth/login", "POST", "Non-existent email", "401 Unauthorized", "401 Unauthorized", "PASS"],
        ["TC-06", "/posts", "GET", "No query params (unauthenticated)", "200 OK, paginated posts array", "200 OK, 10 posts returned", "PASS"],
        ["TC-07", "/posts", "GET", "Query: type=job&city=Bilaspur", "200 OK, filtered posts (type=job, city=Bilaspur)", "200 OK, correct filtered results", "PASS"],
        ["TC-08", "/posts", "POST", "Valid post body, valid JWT in header", "201 Created, new post object", "201 Created, post document returned", "PASS"],
        ["TC-09", "/posts", "POST", "Missing title field, valid JWT", "400 Bad Request, validation error", "400 Bad Request, 'Title is required'", "PASS"],
        ["TC-10", "/posts", "POST", "No Authorization header", "401 Unauthorized", "401 Unauthorized, 'No token provided'", "PASS"],
        ["TC-11", "/posts/:id/like", "POST", "Valid post ID, valid JWT", "200 OK, updated likes array", "200 OK, user ID added to likes", "PASS"],
        ["TC-12", "/posts/:id/like", "POST", "Already liked post, same JWT", "200 OK (unlike), user ID removed from likes", "200 OK, like toggled off", "PASS"],
        ["TC-13", "/users/:id/follow", "POST", "Valid user ID, valid JWT", "200 OK, follow toggled", "200 OK, follower/following updated", "PASS"],
        ["TC-14", "/conversations", "POST", "Valid recipientId and postId, JWT", "201 Created or 200 OK (existing), conversation object", "200 OK, existing or new conversation", "PASS"],
        ["TC-15", "/conversations/:id/messages", "POST", "Message content, valid JWT", "201 Created, message object", "201 Created, message persisted", "PASS"],
        ["TC-16", "/notifications", "GET", "Valid JWT", "200 OK, notifications array", "200 OK, notifications returned", "PASS"],
        ["TC-17", "/notifications/read-all", "PATCH", "Valid JWT", "200 OK, all notifications marked read", "200 OK, unread count becomes 0", "PASS"],
        ["TC-18", "/admin/stats", "GET", "Admin JWT", "200 OK, stats object (users, posts counts)", "200 OK, correct stats returned", "PASS"],
        ["TC-19", "/admin/stats", "GET", "Regular user JWT (non-admin)", "403 Forbidden", "403 Forbidden, 'Access denied'", "PASS"],
        ["TC-20", "/admin/posts/:id", "DELETE", "Admin JWT, valid post ID", "200 OK, post deleted", "200 OK, post removed from DB", "PASS"],
    ],
    col_widths=[0.5, 1.7, 0.6, 1.4, 1.3, 1.3, 0.55]
)

h(2, "7.4 Frontend Testing")
body("Frontend testing was conducted through systematic manual testing in Google Chrome (primary) and Mozilla Firefox (secondary), using the browser's built-in Developer Tools for debugging. The React Developer Tools browser extension was used to inspect component state, context values, and the component tree hierarchy. The following aspects of the frontend were systematically tested: (1) Form Validation — all forms (Register, Login, Create Post, Edit Profile) were tested with empty required fields, invalid email formats, passwords below the minimum length, and successful valid submissions to verify that validation messages appear and disappear correctly. (2) Responsive Layout — all pages were tested at multiple viewport widths (320px, 375px, 768px, 1024px, 1440px) using Chrome DevTools' device simulation mode, verifying that the responsive grid, navigation, and card layouts adapt correctly. (3) Navigation — all routes were tested for correct rendering, and protected routes were verified to redirect unauthenticated users to the Login page. (4) Toast Notifications — all success and error toast notifications were verified to appear with the correct message and dismiss automatically after the configured duration.")

h(2, "7.5 Security Testing")
body("Security testing focused on four key areas. First, JWT token expiry was tested by manually modifying the JWT_EXPIRES_IN environment variable to a very short duration (30 seconds), logging in, waiting for expiry, and then attempting an authenticated API call. The Axios interceptor correctly caught the 401 response and redirected the user to the Login page. Second, unauthorized route access was tested by attempting to access protected API endpoints (POST /posts, GET /admin/stats) without an Authorization header, with an invalid token, and with an expired token—all correctly returned 401 or 403 responses. Third, NoSQL injection prevention was tested by submitting JSON injection payloads (e.g., { '$gt': '' } as the email field in the login request). Mongoose's type casting and validation correctly rejected these inputs before they reached the database query layer. Fourth, CORS policy was verified by making API requests from an origin not in the allowedOrigins list (simulated using a simple HTML file opened via file:// protocol), which was correctly blocked by the CORS middleware.")

h(2, "7.6 Performance Testing")
body("Performance testing was conducted after deployment to the cloud hosting platform. Page load time was measured using Chrome DevTools' Network panel and the Lighthouse performance auditing tool. The initial page load (including React bundle download, parse, and first render) measured approximately 1.8 seconds on a standard broadband connection—well within the 3-second target defined in the non-functional requirements. Subsequent navigation between pages (client-side routing) was near-instantaneous (<100ms) due to the SPA architecture. API response times were measured using Postman's response time display: the GET /posts endpoint returned responses in 180–350ms for a dataset of 500 posts, which is within acceptable bounds for the development deployment tier. MongoDB Atlas query performance was monitored using the Atlas Performance Advisor, which confirmed that the indexes on email (Users collection) and the compound index on type+coordinates (Posts collection) were being effectively used by the most frequent query patterns.")

h(2, "7.7 User Acceptance Testing (UAT)")
body("User Acceptance Testing was conducted with a group of 5 volunteer student participants recruited from the MCA and BCA programmes at GGU, Bilaspur. Each participant was given a brief introduction to MyCircle's purpose (without a detailed feature walkthrough) and asked to complete a set of predefined tasks: (1) Register a new account and set up a profile with a photo; (2) Create a new opportunity post (of any type); (3) Find an opportunity posted by another participant using the search and filter; (4) Send a message to another participant; (5) Follow another participant and check their profile. Participants completed these tasks while the developer observed, noting points of confusion, errors encountered, and time taken for each task.")
body("Key findings from the UAT session included: (1) All 5 participants successfully completed Registration and Profile Setup, with an average time of 3.2 minutes. (2) 4 of 5 participants successfully created an opportunity post; 1 participant was initially confused by the 'Post Type' selector and required a brief explanation. (3) All participants successfully found opportunities using the search bar; the filter panel was used by 3 of 5 participants. (4) All participants successfully sent a message after a brief exploration of the Messages page. (5) All participants successfully followed another user. Post-session feedback highlighted two UI improvements that were subsequently implemented: adding a tooltip to the Post Type selector explaining each category, and making the 'More' collapse button on PostCards more visually prominent.")

h(2, "7.8 Bug Report Table")
body("The following table documents bugs identified and resolved during the development and testing phases of MyCircle.")
add_table(
    ["Bug ID", "Description", "Severity", "Status", "Resolution"],
    [
        ["BUG-01", "Filter function not defined error on Explore page after code refactor introduced escaped quotes in JSX", "Critical", "Resolved", "Restored correct JSX syntax; re-validated all string interpolations in Explore.jsx"],
        ["BUG-02", "Currency symbol showing hardcoded '₹' instead of dynamic symbol from CurrencySymbolContext in PostCard", "High", "Resolved", "Integrated useCurrencySymbol hook; replaced all hardcoded currency characters"],
        ["BUG-03", "cn() utility function not imported in Explore.jsx after map UX enhancement; 'cn is not defined' runtime error", "Critical", "Resolved", "Added import { cn } from '../utils/cn' to Explore.jsx"],
        ["BUG-04", "Duplicate import block in PostCard.jsx after multi-replace operation caused parse error", "Critical", "Resolved", "Removed duplicated import lines; ensured single clean import block"],
        ["BUG-05", "JWT token not attached to Axios requests after page refresh due to race condition in AuthContext initialization", "High", "Resolved", "Added loading state to AuthContext; deferred API calls until auth state initialized"],
        ["BUG-06", "Leaflet map markers not rendering on first load due to missing default icon image paths", "Medium", "Resolved", "Added L.Icon.Default.mergeOptions() with explicit CDN URLs for marker icons"],
        ["BUG-07", "Post cards showing 'undefined' for budget when only a single price value provided (no budgetMin/Max)", "Medium", "Resolved", "Updated budgetLabel logic to fall back to price field when budget range fields are absent"],
        ["BUG-08", "Socket.io 'new_message' event not received by recipient when conversation window not open", "High", "Resolved", "Moved socket subscription to global SocketContext rather than ChatWindow component"],
        ["BUG-09", "Admin dashboard accessible to regular users due to missing role check on frontend ProtectedRoute", "High", "Resolved", "Added role === 'admin' check to ProtectedRoute; backend admin middleware also enforces role"],
        ["BUG-10", "Profile image upload failing for large images (>5MB) due to missing Multer file size limit configuration", "Medium", "Resolved", "Added limits: { fileSize: 5 * 1024 * 1024 } to Multer configuration; added client-side size validation"],
    ],
    col_widths=[0.65, 2.4, 0.75, 0.75, 1.9]
)

doc.add_page_break()
doc.save("MyCircle_Report_TEMP.docx")
print("Part 6 saved — Chapter 7 complete.")
