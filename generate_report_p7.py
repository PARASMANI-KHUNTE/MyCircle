from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document("MyCircle_Report_TEMP.docx")

def h(level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(max(12, 18 - level * 2))
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8 — SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 8: SCREENSHOTS / UI WALKTHROUGH")
body("Note to the student: Replace each [Screenshot Placeholder] section below with an actual screenshot captured from the running MyCircle application. Screenshots should be captured at 1280×800 resolution or higher and inserted as inline images in the final Word document. Each screenshot should be labelled with the figure number and caption as specified.")

h(2, "8.1 Login / Registration Page")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Figure 8.1 — Screenshot Placeholder: Login / Registration Page]")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
body("The Login and Registration page is the first screen that any visitor to MyCircle encounters. The layout features a split-panel design: the left half displays a visually rich hero panel with the MyCircle brand mark, a compelling tagline ('Your campus. Your circle. Your opportunities.'), and an abstract background gradient that reflects the platform's primary colour palette (deep indigo with accent violet tones). The right half contains the authentication form, which toggles between Login and Register modes via a tab interface. The Register form collects Display Name, Email, Password (with a real-time strength indicator bar), College Name, and City. The Login form shows Email and Password fields with a 'Forgot Password' link. Both forms feature a prominent primary-coloured CTA button, inline validation error messages rendered in red below each invalid field, and a loading spinner state that activates while the API call is in progress. The overall visual treatment—rounded input corners, subtle drop shadows, smooth tab transition animation—establishes the premium, modern aesthetic maintained throughout the application.")

h(2, "8.2 Home / Opportunity Feed (Explore Page — List View)")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Figure 8.2 — Screenshot Placeholder: Explore Page — List View]")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
body("The Explore page in List View presents the opportunity feed as a responsive card grid. At the top, a sticky search and filter bar contains a search input with a magnifying glass icon, a row of pill-shaped post type filter buttons (All, Jobs, Services, Sell, Rent), and a Filters toggle button that expands an advanced filter panel. The List/Map view toggle appears at the top right. Below the filter bar, opportunity PostCards are arranged in a 3-column grid on desktop. Each card displays a cover image (or category emoji on a gradient background for image-less posts), a type badge in the top-left of the image, a price badge in the bottom-right of the image, the post title in bold, the author's avatar and name, the posting date and location, and a compact footer with Like, Share, AI Insights, and More/Less toggle buttons. The collapsed card state keeps the feed compact; clicking 'More' smoothly expands the card to reveal Budget Tags, Lifecycle Progress, and (for own posts) Edit/Delete buttons. The visual hierarchy is clean, with the post title as the dominant typographic element and the footer buttons using muted tones that lighten on hover.")

h(2, "8.3 Map View (Explore Page — Map View)")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Figure 8.3 — Screenshot Placeholder: Explore Page — Map View]")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
body("The Map View mode of the Explore page renders an interactive Leaflet map occupying the full content area below the filter bar. The map uses OpenStreetMap tiles and is centred on the user's current GPS location (indicated by a pulsing blue dot with a radiating animation circle). Opportunity posts with GPS coordinates are rendered as custom pin-shaped markers: each marker has a teardrop shape filled with the category's accent colour, an emoji icon representing the post type (briefcase for jobs, wrench for services, tag for sell, house for rent), and a small price badge at the top-right. A glassmorphism-styled floating control panel in the bottom-right corner of the map contains three icon buttons: Zoom In (+), Zoom Out (−), and Recenter (target crosshair icon), plus a Radius Toggle (sparkles icon) that overlays a translucent circle showing the search radius. Clicking a marker opens a card-style popup with the post title, type badge, price, and a 'View Details' link that navigates to the full PostDetails page.")

h(2, "8.4 Student Profile Page")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Figure 8.4 — Screenshot Placeholder: Student Profile Page]")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
body("The Profile page is organized into two sections. The upper hero section spans the full width and features a gradient banner background, the user's circular profile photo (with a camera icon overlay for upload on the own profile), display name in large bold text, college and city below the name, and a row of social metrics (X Followers, X Following, X Posts) that are clickable to open the respective modal lists. For the authenticated user viewing their own profile, an 'Edit Profile' button appears; for other users, a Follow/Unfollow button appears in its place. The reputation badges (Trust Score and Average Rating) appear as small coloured chips below the social metrics. The lower section contains a grid of PostCards displaying all posts authored by the profile user, using the same compact, collapsible card design as the Explore feed. On mobile, the hero section stacks vertically and the card grid collapses to a single column.")

h(2, "8.5 Create / Edit Post Form")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Figure 8.5 — Screenshot Placeholder: Create Post Form]")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
body("The Create Post form page presents a clean, card-contained multi-section form with a sticky header showing the form title ('Create New Post') and a progress indicator. The Post Type selector appears as four visual cards arranged in a 2×2 grid, each with a large emoji, type label, and brief description; the selected type receives a primary-colour border and background highlight. Below the type selector, form fields are organized in logical groups with clear section labels: Basic Info (Title, Description), Pricing (Price, Budget Range with dual sliders for Min and Max, Accepts Barter toggle), Details (Duration in days, Availability, Tags with comma-separated input that renders tags as pill badges), Location (text field plus a compact embedded Leaflet map for coordinate selection), and Images (a drag-and-drop file upload zone that renders selected images as thumbnail previews with remove buttons). The form footer contains a Cancel button and a primary 'Publish Post' CTA button. All validation errors render as red helper text below the relevant field.")

h(2, "8.6 Messaging Interface")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Figure 8.6 — Screenshot Placeholder: Messages Page]")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
body("The Messages page uses a classic two-panel layout on desktop: the left panel lists all conversations, and the right panel shows the active chat thread. The conversations list displays each conversation as a row containing the other participant's avatar, display name, last message preview text, and last message timestamp. Conversations with unread messages are visually highlighted with a bold last-message preview and a small primary-coloured unread count badge. The right panel (ChatWindow) shows the conversation header (participant avatar, name, and a link to their profile), the scrollable message thread, and the message input area. Messages from the authenticated user appear on the right side in primary-coloured rounded bubbles; messages from the other participant appear on the left in card-coloured bubbles. System messages (e.g., 'Request approved — conversation started') appear centred in muted italic text. The message input is a textarea that auto-expands up to three lines, with a Send button (paper plane icon) that activates when the input is non-empty.")

h(2, "8.7 Admin Dashboard")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Figure 8.7 — Screenshot Placeholder: Admin Dashboard]")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
body("The Admin Dashboard is accessible exclusively to administrator accounts and presents a governance-focused interface. The page header clearly identifies the context with an 'Admin Panel' title and shield icon. Below the header, a statistics row displays four metric cards: Total Users (with a trend percentage vs. last week), Total Posts, Total Active Conversations, and Total Notifications Sent. Each metric card uses an icon, a large number in the primary font size, and a subtle background colour differentiated by metric type. Below the stats, two tabbed management tables are presented. The Users tab shows a searchable, paginated table with columns for avatar, user name, email, role badge (Admin or User), registration date, account status (Active/Inactive), and action buttons (Deactivate, Delete). The Posts tab shows a similar table with post title, type badge, author, creation date, status, and view/like counts, with a Delete action button. Each table supports search and sorting by relevant columns.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 9 — LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════
h(1, "CHAPTER 9: LIMITATIONS")

h(2, "9.1 Current Technical Limitations")
body("The most immediately apparent technical limitation of the current version of MyCircle is the absence of a dedicated mobile application. While the web application is fully responsive and functional on mobile browsers, a native or cross-platform mobile app (built using React Native, for example) would provide a significantly superior mobile user experience through device-native UI patterns, push notification support via Firebase Cloud Messaging, offline caching, camera and gallery integration for post image uploads, and access to the device's GPS without requiring browser permission prompts. The current mobile web experience, while functional, cannot fully replicate these capabilities.")
body("A second technical limitation is the recommendation and content discovery algorithm. The current Explore feed displays posts in reverse-chronological order (most recent first), with no personalization based on the user's interests, follow graph, or browsing history. A machine learning-based recommendation engine—using collaborative filtering (recommending posts liked by users similar to the current user) or content-based filtering (recommending posts with tags similar to posts the user has interacted with)—would substantially improve the relevance of the feed. Implementing such a system would require a data pipeline for collecting interaction events, a model training environment, and a serving infrastructure for real-time recommendations, all of which are beyond the scope of the current Minor Project.")
body("The full-text search capability in the current version is limited to basic MongoDB regex pattern matching on the title and description fields. This approach is inefficient for large datasets (regex queries cannot use standard MongoDB indexes) and does not support fuzzy matching, synonym expansion, or relevance-scored results. A production-grade search experience would require integration with a dedicated search engine such as Elasticsearch or MongoDB Atlas Search, which provides inverted-index-based full-text search with relevance ranking, autocomplete, and typo tolerance.")
body("The JWT authentication tokens in the current implementation are stored in the browser's localStorage. While this is a common practice in SPAs, it creates a vulnerability to Cross-Site Scripting (XSS) attacks: if a malicious script is somehow injected into the page, it can read the localStorage token. The more secure alternative—storing the JWT in an HttpOnly cookie that JavaScript cannot access—is identified as a prioritized security enhancement for a future version. Additionally, token refresh logic (automatically issuing a new short-lived token before the current one expires, using a longer-lived refresh token) is not implemented in the current version, meaning users must re-login every 7 days.")

h(2, "9.2 Platform / Deployment Limitations")
body("MyCircle is currently deployed on Render's free tier, which introduces a significant operational constraint: the server process spins down after 15 minutes of inactivity, and the first request after a spin-down incurs a 30–60 second cold start delay while the container reinitializes. For a production platform serving real users, this cold start latency is unacceptable—a user attempting to log in during a cold start would experience a 30+ second wait with no feedback. Resolving this requires upgrading to a paid hosting plan that keeps the server always-on, or implementing a 'keep-alive' ping mechanism (a scheduled request every 10 minutes to prevent spin-down).")
body("The MongoDB Atlas free tier (M0 cluster) provides 512 MB of storage, which is sufficient for development and initial testing but will become a constraint as the platform grows. With an average post document size of approximately 2KB (including embedded image URLs and arrays), the free tier can accommodate approximately 256,000 posts before storage is exhausted. Scaling beyond the free tier requires upgrading to a paid Atlas cluster (M10 or above), which introduces monthly infrastructure costs. Similarly, Cloudinary's free tier provides 25 GB of storage and 25 GB of bandwidth per month—sufficient for a small pilot but limiting for a larger student community.")

h(2, "9.3 Business and Scope Limitations")
body("The current version of MyCircle is designed as a single-community platform: all users can see all posts regardless of their college or city (with geographic filtering available but not enforced). While this open design maximizes content availability during the initial launch phase, it is architecturally incompatible with a multi-college, institution-specific community model where students at College A see a distinct feed from students at College B. Implementing true community isolation would require a community/tenant data model with institution-specific feeds and privacy boundaries, which represents a significant architectural enhancement beyond the current scope.")
body("The platform currently lacks a payment infrastructure, which limits the potential for premium listing features, sponsored opportunities from local businesses, and a subscription model for advanced student profiles. Integrating a payment gateway (such as Razorpay or Stripe) for the Indian market would require PCI-DSS compliance considerations, KYC verification for business accounts, and significant additional development effort. Finally, content moderation in the current version is entirely manual—administrators must proactively review reported posts and user complaints. An automated content moderation layer (using ML-based content classification or keyword filtering) would significantly reduce the administrative burden as the platform scales.")

h(2, "9.4 Security Limitations")
body("As noted in Section 9.1, the JWT storage in localStorage represents a potential XSS vulnerability. While MyCircle's React codebase does not use dangerouslySetInnerHTML or any mechanism that would allow XSS injection in the current implementation, a dependency vulnerability or future code change could introduce this risk. Migrating to HttpOnly cookie-based token storage is a security best practice that should be implemented before large-scale public deployment.")
body("Two-Factor Authentication (2FA) is not implemented in the current version. For a platform dealing with student academic and personal information, 2FA via SMS OTP or TOTP (Time-based One-Time Password, e.g., Google Authenticator) would significantly strengthen account security against credential theft. Email verification upon registration is also absent—users can register with any email address without verifying ownership, which opens the platform to spam account creation and reduces the trustworthiness of user profiles. Both 2FA and email verification are prioritized security enhancements for the next version.")
body("Rate limiting is partially implemented via the Express application but is not comprehensively enforced on all endpoints. The authentication endpoints (login, register) are the most critical targets for rate limiting to prevent brute-force attacks. A production deployment should use a dedicated rate-limiting middleware (express-rate-limit) with a Redis-backed store for accurate rate tracking across multiple server instances, combined with CAPTCHA verification on the registration and login forms after a threshold number of failed attempts.")

doc.add_page_break()
doc.save("MyCircle_Report_TEMP.docx")
print("Part 7 saved — Chapters 8 and 9 complete.")
